"""
Document Processing Service
Handles PDF, DOCX, and email document processing
"""

import logging
import asyncio
import aiohttp
import io
from typing import List, Dict, Any, Optional
from pathlib import Path
import tempfile
import ssl

# Fix SSL certificate verification issues
ssl._create_default_https_context = ssl._create_stdlib_context

# Document processing imports
import PyPDF2
from docx import Document
import email
from email.mime.text import MIMEText

# Add these imports
import re
from nltk.tokenize import sent_tokenize
from sentence_transformers import SentenceTransformer
import numpy as np

from app.core.config import settings

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Document processing and text extraction service"""
    
    def __init__(self):
        """Initialize document processor"""
        self.embedding_model = SentenceTransformer(settings.EMBEDDING_MODEL, device='cpu')
        logger.info(f"Initialized document processor with embedding model: {settings.EMBEDDING_MODEL}")
    
    async def download_document(self, url: str) -> bytes:
        """
        Download document from URL
        
        Args:
            url: Document URL (blob URL, HTTP URL, etc.)
            
        Returns:
            Document content as bytes
        """
        try:
            async with aiohttp.ClientSession() as session:
                async with session.get(url) as response:
                    if response.status == 200:
                        content = await response.read()
                        logger.info(f"Downloaded document from {url}, size: {len(content)} bytes")
                        return content
                    else:
                        raise Exception(f"Failed to download document: HTTP {response.status}")
        except Exception as e:
            logger.error(f"Error downloading document from {url}: {str(e)}")
            raise
    
    def extract_text_from_pdf(self, pdf_content: bytes) -> str:
        """
        Extract text from PDF content
        
        Args:
            pdf_content: PDF file content as bytes
            
        Returns:
            Extracted text
        """
        try:
            pdf_file = io.BytesIO(pdf_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page in pdf_reader.pages:
                try:
                    text += page.extract_text() + "\n"
                except Exception:
                    continue
            
            # Remove redundant whitespace
            text = re.sub(r'\s+', ' ', text).strip()
            logger.info(f"Extracted {len(text)} characters from PDF")
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise
    
    def extract_text_from_docx(self, docx_content: bytes) -> str:
        """
        Extract text from DOCX content
        
        Args:
            docx_content: DOCX file content as bytes
            
        Returns:
            Extracted text
        """
        try:
            docx_file = io.BytesIO(docx_content)
            doc = Document(docx_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            logger.info(f"Extracted {len(text)} characters from DOCX")
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise
    
    def extract_text_from_email(self, email_content: bytes) -> str:
        """
        Extract text from email content
        
        Args:
            email_content: Email content as bytes
            
        Returns:
            Extracted text including headers and body
        """
        try:
            email_str = email_content.decode('utf-8', errors='ignore')
            msg = email.message_from_string(email_str)
            
            # Extract headers
            text = f"Subject: {msg.get('Subject', 'N/A')}\n"
            text += f"From: {msg.get('From', 'N/A')}\n"
            text += f"To: {msg.get('To', 'N/A')}\n"
            text += f"Date: {msg.get('Date', 'N/A')}\n\n"
            
            # Extract body
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        text += part.get_payload(decode=True).decode('utf-8', errors='ignore')
                    elif part.get_content_type() == "text/html":
                        # Basic HTML text extraction (you might want to use BeautifulSoup for better extraction)
                        html_content = part.get_payload(decode=True).decode('utf-8', errors='ignore')
                        # Simple HTML tag removal
                        import re
                        text += re.sub('<[^<]+?>', '', html_content)
            else:
                text += msg.get_payload(decode=True).decode('utf-8', errors='ignore')
            
            logger.info(f"Extracted {len(text)} characters from email")
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from email: {str(e)}")
            raise
    
    def chunk_text(self, text: str, chunk_size: int = None, overlap: int = None) -> List[Dict[str, Any]]:
        """
        Split text into chunks for processing
        
        Args:
            text: Input text to chunk
            chunk_size: Size of each chunk (defaults to settings.CHUNK_SIZE)
            overlap: Overlap between chunks (defaults to settings.CHUNK_OVERLAP)
            
        Returns:
            List of text chunks with metadata
        """
        chunk_size = chunk_size or settings.CHUNK_SIZE
        overlap = overlap or settings.CHUNK_OVERLAP
        
        # Use sentence tokenization for better chunk boundaries
        sentences = sent_tokenize(text)
        chunks = []
        current_chunk = ""
        chunk_id = 0
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) > chunk_size and current_chunk:
                chunks.append({
                    "id": chunk_id,
                    "text": current_chunk.strip(),
                    "length": len(current_chunk),
                })
                chunk_id += 1
                # Start new chunk with overlap
                current_chunk = current_chunk[-overlap:] + " " + sentence if overlap > 0 else sentence
            else:
                current_chunk += " " + sentence
        
        if current_chunk.strip():
            chunks.append({
                "id": chunk_id,
                "text": current_chunk.strip(),
                "length": len(current_chunk),
            })
        
        logger.info(f"Created {len(chunks)} chunks from text")
        return chunks
    
    def create_embeddings(self, texts: List[str]) -> np.ndarray:
        """
        Create embeddings for text chunks
        
        Args:
            texts: List of text strings to embed
            
        Returns:
            Numpy array of embeddings
        """
        try:
            # Batch processing for efficiency
            batch_size = 32
            embeddings = []
            for i in range(0, len(texts), batch_size):
                batch = texts[i:i+batch_size]
                embeddings.append(self.embedding_model.encode(batch, convert_to_numpy=True))
            
            return np.vstack(embeddings)
        except Exception as e:
            logger.error(f"Error creating embeddings: {str(e)}")
            raise
    
    async def process_document(self, document_url: str) -> Dict[str, Any]:
        """
        Process document from URL - download, extract text, chunk, and embed
        
        Args:
            document_url: URL of the document to process
            
        Returns:
            Processed document data with chunks and embeddings
        """
        try:
            # Download document
            content = await self.download_document(document_url)
            
            # Determine file type and extract text
            text = ""
            file_type = "unknown"
            
            # Improved file type detection
            if content.startswith(b'%PDF'):
                try:
                    text = self.extract_text_from_pdf(content)
                    file_type = "pdf"
                except Exception as pdf_err:
                    logger.error(f"PDF extraction failed: {pdf_err}")
                    text = content.decode('utf-8', errors='ignore')
                    file_type = "text"
            elif content[:2] == b'PK':
                try:
                    text = self.extract_text_from_docx(content)
                    file_type = "docx"
                except Exception as docx_err:
                    logger.error(f"DOCX extraction failed: {docx_err}")
                    text = content.decode('utf-8', errors='ignore')
                    file_type = "text"
            elif b'From:' in content[:1000] or b'Subject:' in content[:1000]:
                text = self.extract_text_from_email(content)
                file_type = "email"
            else:
                # Try to decode as plain text
                text = content.decode('utf-8', errors='ignore')
                file_type = "text"

            # If text extraction failed, raise a clear error
            if not text or text.strip() == "":
                logger.error(f"Text extraction failed for document: {document_url}. No text could be extracted.")
                raise Exception("Text extraction failed. File may be corrupted or unsupported format.")
            
            # Create chunks
            chunks = self.chunk_text(text)
            
            # Create embeddings
            chunk_texts = [chunk["text"] for chunk in chunks]
            embeddings = self.create_embeddings(chunk_texts)
            
            # Add embeddings to chunks
            for i, chunk in enumerate(chunks):
                chunk["embedding"] = embeddings[i].tolist()
            
            processed_doc = {
                "url": document_url,
                "file_type": file_type,
                "total_length": len(text),
                "num_chunks": len(chunks),
                "chunks": chunks,
                "full_text": text[:5000],  # Store first 5000 chars for reference
            }
            
            logger.info(f"Successfully processed document: {file_type}, {len(chunks)} chunks")
            return processed_doc
            
        except Exception as e:
            logger.error(f"Error processing document {document_url}: {str(e)}")
            raise